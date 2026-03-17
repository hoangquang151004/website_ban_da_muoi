from __future__ import annotations

import hashlib
import logging
from datetime import datetime, timezone
from pathlib import Path

from langchain_core.documents import Document
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from app.db.session import AsyncSessionLocal
from app.models.data_source import (
    DataChunk,
    DataSource,
    DataSourceStatus,
    IndexJob,
    IndexJobStatus,
)
from app.services.ai_agent.vector_store import get_vector_store

logger = logging.getLogger(__name__)


def parse_source_file(file_path: str) -> str:
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"Source file not found: {file_path}")

    suffix = path.suffix.lower()
    raw = path.read_bytes()

    if suffix in {".txt", ".md", ".csv", ".json", ".log", ".py"}:
        return raw.decode("utf-8", errors="ignore")

    if suffix == ".pdf":
        try:
            from pypdf import PdfReader

            reader = PdfReader(file_path)
            return "\n".join((page.extract_text() or "") for page in reader.pages)
        except Exception:
            return raw.decode("utf-8", errors="ignore")

    if suffix == ".docx":
        try:
            from docx import Document as DocxDocument

            doc = DocxDocument(file_path)
            return "\n".join(p.text for p in doc.paragraphs)
        except Exception:
            return raw.decode("utf-8", errors="ignore")

    return raw.decode("utf-8", errors="ignore")


def chunk_text(text: str, chunk_size: int = 1200, overlap: int = 150) -> list[str]:
    normalized = "\n".join(line.strip() for line in text.splitlines() if line.strip())
    if not normalized:
        return []

    step = max(1, chunk_size - overlap)
    chunks: list[str] = []
    for start in range(0, len(normalized), step):
        piece = normalized[start : start + chunk_size]
        if piece:
            chunks.append(piece)
    return chunks


def text_hash(text: str) -> str:
    return hashlib.sha256(text.encode("utf-8")).hexdigest()


def chroma_chunk_id(source_id: str, chunk_index: int, version: int) -> str:
    return f"{source_id}:{chunk_index}:{version}"


def _is_chroma_present(vs, chunk_id: str) -> bool:
    collection = getattr(vs, "_collection", None)
    if collection is None:
        return False
    found = collection.get(ids=[chunk_id], include=[])
    raw_ids = found.get("ids")
    ids = list(raw_ids) if raw_ids is not None else []
    return chunk_id in ids


def _copy_embedding_from_cached_chunk(vs, cached_id: str, new_id: str, document: str, metadata: dict) -> bool:
    collection = getattr(vs, "_collection", None)
    if collection is None:
        return False

    cached = collection.get(ids=[cached_id], include=["embeddings"])
    raw_ids = cached.get("ids")
    ids = list(raw_ids) if raw_ids is not None else []
    raw_embeddings = cached.get("embeddings")
    embeddings = list(raw_embeddings) if raw_embeddings is not None else []
    if cached_id not in ids or len(embeddings) == 0:
        return False

    embedding = embeddings[0]
    collection.upsert(
        ids=[new_id],
        embeddings=[embedding],
        documents=[document],
        metadatas=[metadata],
    )
    return True


async def validate_source_consistency(
    db: AsyncSession,
    source_id: str,
    version: int,
) -> dict:
    vs = get_vector_store()
    collection = getattr(vs, "_collection", None)
    if collection is None:
        return {"ok": True, "expected": 0, "missing": []}

    result = await db.execute(
        select(DataChunk.chroma_id).where(
            DataChunk.source_id == source_id,
            DataChunk.version == version,
        )
    )
    expected_ids = [row[0] for row in result.all() if row[0]]
    if not expected_ids:
        return {"ok": True, "expected": 0, "missing": []}

    existing = collection.get(ids=expected_ids, include=[])
    raw_found_ids = existing.get("ids")
    found_ids = set(list(raw_found_ids) if raw_found_ids is not None else [])
    missing = [cid for cid in expected_ids if cid not in found_ids]

    return {"ok": len(missing) == 0, "expected": len(expected_ids), "missing": missing}


async def _upsert_data_chunk_record(
    db: AsyncSession,
    source_id: str,
    chunk_index: int,
    version: int,
    t_hash: str,
    c_id: str,
) -> None:
    result = await db.execute(
        select(DataChunk).where(
            DataChunk.source_id == source_id,
            DataChunk.chunk_index == chunk_index,
            DataChunk.version == version,
        )
    )
    row = result.scalar_one_or_none()
    if row is None:
        db.add(
            DataChunk(
                source_id=source_id,
                chunk_index=chunk_index,
                version=version,
                text_hash=t_hash,
                chroma_id=c_id,
            )
        )
    else:
        row.text_hash = t_hash
        row.chroma_id = c_id


async def ingest_source_chunks(
    db: AsyncSession,
    source: DataSource,
    version: int,
    chunks: list[str],
) -> dict:
    now_iso = datetime.now(timezone.utc).isoformat()
    vs = get_vector_store()

    docs_to_embed: list[Document] = []
    ids_to_embed: list[str] = []
    cache_hits = 0
    embedded = 0

    for idx, piece in enumerate(chunks):
        c_id = chroma_chunk_id(source.id, idx, version)
        t_hash = text_hash(piece)
        metadata = {
            "source_id": source.id,
            "source_name": source.name,
            "category": source.category,
            "chunk_index": idx,
            "text_hash": t_hash,
            "version": version,
            "created_at": now_iso,
        }

        await _upsert_data_chunk_record(
            db=db,
            source_id=source.id,
            chunk_index=idx,
            version=version,
            t_hash=t_hash,
            c_id=c_id,
        )

        # Idempotent rerun for same source/chunk/version.
        if _is_chroma_present(vs, c_id):
            continue

        # Reuse existing embedding by text_hash if possible.
        cached_result = await db.execute(
            select(DataChunk.chroma_id)
            .where(DataChunk.text_hash == t_hash, DataChunk.chroma_id.is_not(None))
            .order_by(DataChunk.id.desc())
            .limit(1)
        )
        cached_row = cached_result.first()
        cached_id = cached_row[0] if cached_row else None
        if cached_id and _copy_embedding_from_cached_chunk(vs, cached_id, c_id, piece, metadata):
            cache_hits += 1
            continue

        docs_to_embed.append(Document(page_content=piece, metadata=metadata))
        ids_to_embed.append(c_id)

    if docs_to_embed:
        vs.add_documents(docs_to_embed, ids=ids_to_embed)
        embedded = len(docs_to_embed)

    return {
        "total_chunks": len(chunks),
        "embedded_chunks": embedded,
        "cache_hits": cache_hits,
    }


async def _set_job_state(
    db: AsyncSession,
    job: IndexJob,
    source: DataSource | None,
    *,
    status: IndexJobStatus,
    source_status: DataSourceStatus,
    progress: int | None = None,
    error: str | None = None,
    started_at: datetime | None = None,
    finished_at: datetime | None = None,
) -> None:
    job.status = status
    if source is not None:
        source.status = source_status
    if progress is not None:
        job.progress = progress
    job.error = error
    if started_at is not None:
        job.started_at = started_at
    if finished_at is not None:
        job.finished_at = finished_at
    await db.flush()


async def run_index_job(job_id: str) -> None:
    logger.info("Index worker start job_id=%s", job_id)
    async with AsyncSessionLocal() as db:
        try:
            job_result = await db.execute(select(IndexJob).where(IndexJob.id == job_id))
            job = job_result.scalar_one_or_none()
            if job is None:
                logger.warning("Index worker skip missing job_id=%s", job_id)
                return

            source_result = await db.execute(select(DataSource).where(DataSource.id == job.source_id))
            source = source_result.scalar_one_or_none()
            if source is None or source.is_deleted:
                await _set_job_state(
                    db,
                    job,
                    source,
                    status=IndexJobStatus.failed,
                    source_status=DataSourceStatus.failed if source else DataSourceStatus.deleted,
                    progress=100,
                    error="Source not found or deleted",
                    finished_at=datetime.now(timezone.utc),
                )
                await db.commit()
                return

            await _set_job_state(
                db,
                job,
                source,
                status=IndexJobStatus.processing,
                source_status=DataSourceStatus.processing,
                progress=5,
                error=None,
                started_at=datetime.now(timezone.utc),
            )

            text = parse_source_file(source.file_path or "")
            await _set_job_state(
                db,
                job,
                source,
                status=IndexJobStatus.processing,
                source_status=DataSourceStatus.processing,
                progress=20,
            )

            chunks = chunk_text(text)
            if not chunks:
                raise RuntimeError("No readable content extracted from source")

            await _set_job_state(
                db,
                job,
                source,
                status=IndexJobStatus.processing,
                source_status=DataSourceStatus.processing,
                progress=45,
            )

            stats = await ingest_source_chunks(
                db=db,
                source=source,
                version=source.current_version,
                chunks=chunks,
            )

            await _set_job_state(
                db,
                job,
                source,
                status=IndexJobStatus.processing,
                source_status=DataSourceStatus.processing,
                progress=90,
            )

            consistency = await validate_source_consistency(
                db=db,
                source_id=source.id,
                version=source.current_version,
            )
            if not consistency["ok"]:
                raise RuntimeError(
                    f"Chroma consistency mismatch, missing={len(consistency['missing'])}"
                )

            await _set_job_state(
                db,
                job,
                source,
                status=IndexJobStatus.indexed,
                source_status=DataSourceStatus.indexed,
                progress=100,
                error=None,
                finished_at=datetime.now(timezone.utc),
            )
            await db.commit()
            logger.info(
                "Index worker done job_id=%s total_chunks=%s embedded=%s cache_hits=%s",
                job_id,
                stats["total_chunks"],
                stats["embedded_chunks"],
                stats["cache_hits"],
            )
        except Exception as exc:
            logger.exception("Index worker failed job_id=%s error=%s", job_id, exc)
            try:
                job_result = await db.execute(select(IndexJob).where(IndexJob.id == job_id))
                job = job_result.scalar_one_or_none()
                if job is not None:
                    source_result = await db.execute(select(DataSource).where(DataSource.id == job.source_id))
                    source = source_result.scalar_one_or_none()
                    if source is not None:
                        await _set_job_state(
                            db,
                            job,
                            source,
                            status=IndexJobStatus.failed,
                            source_status=DataSourceStatus.failed,
                            progress=100,
                            error=str(exc)[:2000],
                            finished_at=datetime.now(timezone.utc),
                        )
                        await db.commit()
            except Exception:
                await db.rollback()
